from google.adk.agents.llm_agent import Agent
from google.adk.tools import ToolContext
from google.adk.tools import load_artifacts
from google.genai import types
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


async def create_docx(
    proposal_markdown: str,
    image_filenames: list,
    output_filename: str,
    tool_context: 'ToolContext'
):
    """
    Creates a professional DOCX document from markdown text and images from session artifacts.
    
    Uses ADK's shared artifact context - images saved by visual_generator 
    in the same session are automatically available here.
    
    Args:
        proposal_markdown: Full proposal text in Markdown format
        image_filenames: List of image filenames to include (e.g., ['investment_breakdown.png'])
        output_filename: Name of the output DOCX file (must end with .docx)
    """
    logger.info("=" * 80)
    logger.info("🟢🟢🟢 DOCX ASSEMBLER FUNCTION STARTED 🟢🟢🟢")
    logger.info(f"Markdown length: {len(proposal_markdown)} characters")
    logger.info(f"Images to include: {image_filenames}")
    logger.info(f"Output filename: {output_filename}")
    logger.info("=" * 80)
    
    try:
        # Create document
        logger.info("Step 1: Creating Document object...")
        doc = Document()
        
        # === STYLING ===
        COMARCH_BLUE = RGBColor(31, 60, 136)  # #1F3C88
        
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        logger.info("✓ Document styling applied")
        
        # === PARSE MARKDOWN ===
        logger.info("Step 2: Parsing markdown...")
        lines = proposal_markdown.split('\n')
        
        in_table = False
        table_lines = []
        sections_added = 0
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if not in_table:
                    doc.add_paragraph()
                continue
            
            # Headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:].strip(), level=1)
                heading.runs[0].font.color.rgb = COMARCH_BLUE
                sections_added += 1
                continue
            
            if line.startswith('## '):
                heading = doc.add_heading(line[3:].strip(), level=2)
                heading.runs[0].font.color.rgb = COMARCH_BLUE
                sections_added += 1
                continue
            
            if line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
                continue
            
            # Tables
            if '|' in line:
                if not in_table:
                    in_table = True
                    table_lines = [line]
                else:
                    table_lines.append(line)
                continue
            else:
                if in_table:
                    _add_markdown_table(doc, table_lines, COMARCH_BLUE)
                    table_lines = []
                    in_table = False
            
            # Lists
            if line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
                continue
            
            if re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line).strip()
                doc.add_paragraph(text, style='List Number')
                continue
            
            # Regular paragraph
            doc.add_paragraph(line)
        
        logger.info(f"✓ Markdown parsed: {sections_added} sections")

        if in_table and table_lines:
            _add_markdown_table(doc, table_lines, COMARCH_BLUE)

        # === INSERT IMAGES (ADK ARTIFACT-FIRST APPROACH) ===
        logger.info("Step 3: Loading images from session artifacts...")
        
        artifact_names = await tool_context.list_artifacts()
        logger.info(f"Found {len(artifact_names)} artifacts in session: {list(artifact_names)}")
        
        valid_images = []
        
        for artifact_name in artifact_names:
            try:
                artifact_name_str = str(artifact_name)
                
                if output_filename in artifact_name_str:
                    logger.info(f"Skipping output file: {artifact_name_str}")
                    continue
                
                is_image_file = any(ext in artifact_name_str.lower() for ext in ['.png', '.jpg', '.jpeg', '.gif'])
                
                if not is_image_file:
                    logger.info(f"Skipping non-image artifact: {artifact_name_str}")
                    continue
                
                logger.info(f"Loading image artifact: {artifact_name_str}")
                
                artifact_part = await tool_context.load_artifact(artifact_name_str)
                
                if not artifact_part:
                    logger.warning(f"⚠ Could not load artifact: {artifact_name_str}")
                    continue
                
                logger.info(f"Artifact loaded, type: {type(artifact_part)}")
                
                image_bytes = None
                
                if hasattr(artifact_part, 'inline_data') and artifact_part.inline_data:
                    if hasattr(artifact_part.inline_data, 'data'):
                        image_bytes = artifact_part.inline_data.data
                        mime_type = getattr(artifact_part.inline_data, 'mime_type', 'unknown')
                        logger.info(f"✓ Extracted from inline_data: {len(image_bytes)} bytes, mime: {mime_type}")
                elif hasattr(artifact_part, 'data'):
                    image_bytes = artifact_part.data
                    logger.info(f"✓ Extracted from data: {len(image_bytes)} bytes")
                elif isinstance(artifact_part, bytes):
                    image_bytes = artifact_part
                    logger.info(f"✓ Direct bytes: {len(image_bytes)} bytes")
                else:
                    logger.warning(f"⚠ Unknown artifact structure: {type(artifact_part)}, attrs: {dir(artifact_part)}")
                
                if image_bytes:
                    clean_name = artifact_name_str.replace('user:', '').replace('user_', '').replace('.png', '').replace('_', ' ').title()
                    valid_images.append({
                        'name': clean_name,
                        'bytes': image_bytes,
                        'original_key': artifact_name_str
                    })
                    logger.info(f"✓ Added to queue: {clean_name} ({len(image_bytes)} bytes)")
                
            except Exception as load_err:
                logger.error(f"❌ Error loading artifact {artifact_name}: {load_err}")
                import traceback
                traceback.print_exc()
        
        images_inserted = 0
        if valid_images:
            doc.add_page_break()
            heading = doc.add_heading('Visual Analysis & Charts', level=1)
            heading.runs[0].font.color.rgb = COMARCH_BLUE
            
            for img_data in valid_images:
                try:
                    doc.add_heading(img_data['name'], level=2)
                    
                    image_stream = io.BytesIO(img_data['bytes'])
                    doc.add_picture(image_stream, width=Inches(6))
                    doc.add_paragraph()
                    
                    images_inserted += 1
                    logger.info(f"✓ Inserted image: {img_data['name']}")
                except Exception as insert_err:
                    logger.error(f"❌ Failed to insert image {img_data['name']}: {insert_err}")
                    p = doc.add_paragraph(f"[Error inserting chart: {img_data['name']}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            logger.info(f"✓ Successfully inserted {images_inserted}/{len(valid_images)} images")
        elif image_filenames:
            logger.warning(f"⚠ No image artifacts found in session. Adding placeholders for: {image_filenames}")
            doc.add_page_break()
            heading = doc.add_heading('Visual Analysis & Charts', level=1)
            heading.runs[0].font.color.rgb = COMARCH_BLUE
            
            for filename in image_filenames:
                clean_name = str(filename).replace('user:', '').replace('user_', '').replace('.png', '').replace('_', ' ').title()
                doc.add_heading(clean_name, level=2)
                p = doc.add_paragraph()
                p.add_run(f'[Chart: {filename} - image not found in session artifacts]').italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        else:
            logger.info("Step 3: No images to insert")
        
        # === SAVE DOCUMENT ===
        logger.info("Step 4: Saving document to buffer...")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()
        
        logger.info(f"✓ Document saved to buffer: {len(docx_bytes)} bytes")
        
        # === SAVE TO ARTIFACT STORAGE ===
        logger.info(f"Step 5: Saving to artifact storage as 'user:{output_filename}'...")
        
        # Create Part object
        docx_part = types.Part.from_bytes(
            data=docx_bytes,
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Save artifact
        await tool_context.save_artifact(
            f'user:{output_filename}',
            docx_part
        )
        
        logger.info("✓ Artifact saved to storage")
        
        # Verify artifact was saved
        logger.info("Step 6: Verifying artifact...")
        artifacts = await tool_context.list_artifacts()
        docx_found = False
        for artifact in artifacts:
            artifact_name = artifact if isinstance(artifact, str) else (artifact.name if hasattr(artifact, 'name') else str(artifact))
            
            if output_filename in artifact_name:
                docx_found = True
                logger.info(f"✓ VERIFIED: Artifact exists - {artifact_name}")
                break
        
        if not docx_found:
            logger.error(f"❌ WARNING: Could not verify artifact '{output_filename}' in storage!")
        
        logger.info("=" * 80)
        logger.info("🎉🎉🎉 DOCX GENERATION COMPLETED SUCCESSFULLY 🎉🎉🎉")
        logger.info(f"✓ File: {output_filename}")
        logger.info(f"✓ Size: {len(docx_bytes)} bytes")
        logger.info(f"✓ Sections: {sections_added}")
        logger.info(f"✓ Images: {images_inserted}")
        logger.info("=" * 80)
        
        return {
            'status': 'success',
            'filename': output_filename,
            'size_bytes': len(docx_bytes),
            'sections': sections_added,
            'images_inserted': images_inserted,
            'verified': docx_found,
            'detail': f'Professional proposal document created: {output_filename}',
            'message': f'✓ DOCX file saved as {output_filename} ({len(docx_bytes)} bytes) and verified in artifact storage'
        }
        
    except Exception as e:
        logger.error("=" * 80)
        logger.error("❌❌❌ DOCX GENERATION FAILED ❌❌❌")
        logger.error(f"Error: {str(e)}")
        logger.error("=" * 80)
        import traceback
        traceback.print_exc()
        
        return {
            'status': 'failed',
            'error': str(e),
            'message': f'Failed to create DOCX: {str(e)}'
        }


def _add_markdown_table(doc, table_lines, header_color):
    """Helper function to convert Markdown table to Word table"""
    
    # Filter out separator lines
    data_lines = [
        line for line in table_lines 
        if not re.match(r'^\|[\s\-:]+\|$', line)
    ]
    
    if len(data_lines) < 1:
        return
    
    # Parse rows
    rows = []
    for line in data_lines:
        line = line.strip('|').strip()
        cells = [cell.strip() for cell in line.split('|')]
        rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill data
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_data
                
                # Header styling
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = header_color
    
    doc.add_paragraph()


# === AGENT DEFINITION ===

instruction = """You are a **Document Assembly Specialist** for B2B sales proposals.

**Your Task:**
Create professional DOCX documents by combining markdown text with images from the session artifacts.

**How It Works (ADK Artifact System):**
- Images generated by `visual_generator` are automatically saved to the session artifacts
- You can access these images through your `create_docx` tool
- The tool will scan session artifacts and include all images it finds

**How to Call create_docx:**
```
create_docx(
    proposal_markdown="# Proposal\\n\\n## Executive Summary\\n...",
    image_filenames=["investment_breakdown.png", "value_proposition.png"],
    output_filename="Comarch_Sales_Proposal.docx"
)
```

**Parameters:**
- `proposal_markdown`: The full proposal text in Markdown format
- `image_filenames`: List of expected image filenames (for reference/fallback)
- `output_filename`: Name for the output file (e.g., "Comarch_Sales_Proposal.docx")

**Important:**
- Images are loaded from session artifacts automatically
- The tool will find any .png/.jpg files saved by other agents in this session
- You don't need to pass image data - just the filenames for reference

**Final confirmation:**
After successful execution, respond with:
- The DOCX file name that was created
- The file size in bytes
- Number of images inserted
- Confirmation that the file is ready for download
"""



docx_assembler_agent = Agent(
    model='gemini-2.5-flash',
    name='docx_assembler',
    description="Document Assembly Specialist - combines markdown text and images into professional DOCX documents",
    instruction=instruction,
    tools=[create_docx, load_artifacts],
)


root_agent = docx_assembler_agent